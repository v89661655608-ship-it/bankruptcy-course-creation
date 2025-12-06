import json
import base64
import io
from typing import Dict, Any
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.units import cm
except ImportError:
    pass

def handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    '''
    Генерация заявления о признании гражданина банкротом на основе собранных данных
    
    Получает: персональные данные, кредитную историю, доходы, имущество
    Возвращает: сформированное заявление в формате DOCX/PDF
    '''
    method: str = event.get('httpMethod', 'GET')
    
    if method == 'OPTIONS':
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Methods': 'GET, POST, OPTIONS',
                'Access-Control-Allow-Headers': 'Content-Type',
                'Access-Control-Max-Age': '86400'
            },
            'body': '',
            'isBase64Encoded': False
        }
    
    if method == 'POST':
        body = event.get('body', '{}')
        if not body:
            body = '{}'
        body_data = json.loads(body)
        
        personal_data = body_data.get('personalData')
        credit_data = body_data.get('creditData')
        income_data = body_data.get('incomeData')
        property_data = body_data.get('propertyData')
        additional_fields = body_data.get('additionalFields', {})
        benefits_data = body_data.get('benefitsData', {})
        children_data = body_data.get('childrenData', {})
        transactions_data = body_data.get('transactionsData', {})
        debt_reason_data = body_data.get('debtReasonData', {})
        appendices_data = body_data.get('appendicesData', {})
        attachment_motion_data = body_data.get('attachmentMotionData', {})
        absence_motion_data = body_data.get('absenceMotionData', {})
        property_exclusion_motion_data = body_data.get('propertyExclusionMotionData', {})
        debt_discharge_motion_data = body_data.get('debtDischargeMotionData', {})
        doc_format = body_data.get('format', 'docx')
        
        if doc_format == 'attachment-motion':
            if not personal_data:
                return {
                    'statusCode': 400,
                    'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
                    'body': json.dumps({'error': 'Недостаточно данных для генерации ходатайства'}),
                    'isBase64Encoded': False
                }
            docx_base64 = generate_attachment_motion_document(personal_data, additional_fields, attachment_motion_data)
            result = {
                'success': True,
                'document': {
                    'data': docx_base64,
                    'createdAt': datetime.now().isoformat(),
                    'format': 'docx',
                    'fileName': f"ходатайство_приобщение_{personal_data.get('inn', 'doc')}.docx"
                }
            }
            return {
                'statusCode': 200,
                'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
                'body': json.dumps(result),
                'isBase64Encoded': False
            }
        
        if doc_format == 'absence-motion':
            if not personal_data:
                return {
                    'statusCode': 400,
                    'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
                    'body': json.dumps({'error': 'Недостаточно данных для генерации ходатайства'}),
                    'isBase64Encoded': False
                }
            docx_base64 = generate_absence_motion_document(personal_data, additional_fields, absence_motion_data)
            result = {
                'success': True,
                'document': {
                    'data': docx_base64,
                    'createdAt': datetime.now().isoformat(),
                    'format': 'docx',
                    'fileName': f"ходатайство_отсутствие_{personal_data.get('inn', 'doc')}.docx"
                }
            }
            return {
                'statusCode': 200,
                'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
                'body': json.dumps(result),
                'isBase64Encoded': False
            }
        
        if doc_format == 'property-exclusion-motion':
            if not personal_data:
                return {
                    'statusCode': 400,
                    'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
                    'body': json.dumps({'error': 'Недостаточно данных для генерации ходатайства'}),
                    'isBase64Encoded': False
                }
            docx_base64 = generate_property_exclusion_motion_document(personal_data, additional_fields, property_data, property_exclusion_motion_data)
            result = {
                'success': True,
                'document': {
                    'data': docx_base64,
                    'createdAt': datetime.now().isoformat(),
                    'format': 'docx',
                    'fileName': f"ходатайство_исключение_имущества_{personal_data.get('inn', 'doc')}.docx"
                }
            }
            return {
                'statusCode': 200,
                'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
                'body': json.dumps(result),
                'isBase64Encoded': False
            }
        
        if doc_format == 'debt-discharge-motion':
            if not personal_data:
                return {
                    'statusCode': 400,
                    'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
                    'body': json.dumps({'error': 'Недостаточно данных для генерации ходатайства'}),
                    'isBase64Encoded': False
                }
            docx_base64 = generate_debt_discharge_motion_document(personal_data, additional_fields, children_data, debt_discharge_motion_data)
            result = {
                'success': True,
                'document': {
                    'data': docx_base64,
                    'createdAt': datetime.now().isoformat(),
                    'format': 'docx',
                    'fileName': f"ходатайство_освобождение_от_долгов_{personal_data.get('inn', 'doc')}.docx"
                }
            }
            return {
                'statusCode': 200,
                'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
                'body': json.dumps(result),
                'isBase64Encoded': False
            }
        
        if not personal_data or not credit_data:
            return {
                'statusCode': 400,
                'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
                'body': json.dumps({'error': 'Недостаточно данных для генерации заявления'}),
                'isBase64Encoded': False
            }
        
        if doc_format == 'creditors-list':
            docx_base64 = generate_creditors_list_document(
                personal_data, credit_data
            )
            result = {
                'success': True,
                'document': {
                    'data': docx_base64,
                    'createdAt': datetime.now().isoformat(),
                    'format': 'docx',
                    'fileName': f"приложение1_список_кредиторов_{personal_data.get('inn', 'doc')}.docx"
                }
            }
        elif doc_format == 'property-list':
            docx_base64 = generate_property_list_document(
                personal_data, property_data
            )
            result = {
                'success': True,
                'document': {
                    'data': docx_base64,
                    'createdAt': datetime.now().isoformat(),
                    'format': 'docx',
                    'fileName': f"приложение2_опись_имущества_{personal_data.get('inn', 'doc')}.docx"
                }
            }
        elif doc_format == 'pdf':
            pdf_base64 = generate_pdf_document(
                personal_data, credit_data, income_data, property_data, additional_fields, benefits_data, children_data, transactions_data, debt_reason_data
            )
            result = {
                'success': True,
                'document': {
                    'data': pdf_base64,
                    'createdAt': datetime.now().isoformat(),
                    'format': 'pdf',
                    'fileName': f"заявление_банкротство_{personal_data.get('inn', 'doc')}.pdf"
                }
            }
        elif doc_format == 'docx':
            docx_base64 = generate_docx_document(
                personal_data, credit_data, income_data, property_data, additional_fields, benefits_data, children_data, transactions_data, debt_reason_data, appendices_data
            )
            result = {
                'success': True,
                'document': {
                    'data': docx_base64,
                    'createdAt': datetime.now().isoformat(),
                    'format': 'docx',
                    'fileName': f"заявление_банкротство_{personal_data.get('inn', 'doc')}.docx"
                }
            }
        else:
            document_text = generate_bankruptcy_application(
                personal_data, credit_data, income_data, property_data
            )
            result = {
                'success': True,
                'document': {
                    'text': document_text,
                    'createdAt': datetime.now().isoformat(),
                    'format': 'text',
                    'fileName': f"заявление_банкротство_{personal_data.get('inn', 'doc')}.txt"
                }
            }
        
        return {
            'statusCode': 200,
            'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
            'body': json.dumps(result),
            'isBase64Encoded': False
        }
    
    return {
        'statusCode': 405,
        'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
        'body': json.dumps({'error': 'Метод не поддерживается'}),
        'isBase64Encoded': False
    }


def format_number(num):
    """Форматирует число с разделителями тысяч"""
    if not num:
        return "0"
    return f"{float(num):,.0f}".replace(',', ' ')


def get_page_word(pages_str: str) -> str:
    """Склонение слова 'лист' в зависимости от числа"""
    # Извлекаем число из строки
    import re
    numbers = re.findall(r'\d+', pages_str)
    if not numbers:
        return pages_str
    
    num = int(numbers[0])
    
    # Правила склонения
    if num % 10 == 1 and num % 100 != 11:
        word = "лист"
    elif num % 10 in [2, 3, 4] and num % 100 not in [12, 13, 14]:
        word = "листа"
    else:
        word = "листов"
    
    # Если в строке уже есть слово "лист", заменяем его
    if 'лист' in pages_str.lower():
        return pages_str
    
    # Добавляем слово к числу
    return f"{pages_str} {word}"


def decline_full_name_genitive(full_name: str) -> str:
    """Склонение ФИО в родительный падеж (простая реализация)"""
    # Простое склонение для распространенных окончаний
    # Для более точного склонения нужна библиотека pymorphy2
    parts = full_name.split()
    if len(parts) != 3:
        return full_name
    
    surname, name, patronymic = parts
    
    # Склонение фамилии
    if surname.endswith('ова') or surname.endswith('ева') or surname.endswith('ина') or surname.endswith('ына'):
        # Женские фамилии: Петрова -> Петровой, Иванова -> Ивановой
        surname_declined = surname[:-1] + 'ой'
    elif surname.endswith('ая'):
        surname_declined = surname[:-2] + 'ой'
    elif surname.endswith('ов') or surname.endswith('ев') or surname.endswith('ин') or surname.endswith('ын'):
        # Мужские фамилии: Петров -> Петрова, Иванов -> Иванова
        surname_declined = surname + 'а'
    elif surname.endswith('ский') or surname.endswith('цкий'):
        surname_declined = surname[:-2] + 'ого'
    elif surname.endswith('а'):
        surname_declined = surname[:-1] + 'ы'
    else:
        surname_declined = surname + 'а'
    
    # Склонение имени
    if name.endswith('а') or name.endswith('я'):
        name_declined = name[:-1] + 'ы'
    elif name.endswith('й'):
        name_declined = name[:-1] + 'я'
    else:
        name_declined = name + 'а'
    
    # Склонение отчества
    if patronymic.endswith('ич'):
        patronymic_declined = patronymic + 'а'
    elif patronymic.endswith('на'):
        patronymic_declined = patronymic[:-1] + 'ы'
    elif patronymic.endswith('вна') or patronymic.endswith('ична'):
        patronymic_declined = patronymic[:-1] + 'ы'
    else:
        patronymic_declined = patronymic + 'а'
    
    return f"{surname_declined} {name_declined} {patronymic_declined}"


def decline_full_name_accusative(full_name: str) -> str:
    """Склонение ФИО в винительный падеж (кого? что?)"""
    parts = full_name.split()
    if len(parts) != 3:
        return full_name
    
    surname, name, patronymic = parts
    
    # Склонение фамилии
    if surname.endswith('ова') or surname.endswith('ева') or surname.endswith('ина') or surname.endswith('ына'):
        # Женские фамилии: Петрова -> Петрову
        surname_declined = surname[:-1] + 'у'
    elif surname.endswith('ая'):
        surname_declined = surname[:-2] + 'ую'
    elif surname.endswith('ов') or surname.endswith('ев') or surname.endswith('ин') or surname.endswith('ын'):
        # Мужские фамилии: Петров -> Петрова
        surname_declined = surname + 'а'
    elif surname.endswith('ский') or surname.endswith('цкий'):
        surname_declined = surname[:-2] + 'ого'
    elif surname.endswith('а'):
        surname_declined = surname[:-1] + 'у'
    else:
        surname_declined = surname + 'а'
    
    # Склонение имени
    if name.endswith('а') or name.endswith('я'):
        name_declined = name[:-1] + 'у'
    elif name.endswith('й'):
        name_declined = name[:-1] + 'я'
    else:
        name_declined = name + 'а'
    
    # Склонение отчества
    if patronymic.endswith('ич'):
        patronymic_declined = patronymic + 'а'
    elif patronymic.endswith('на'):
        patronymic_declined = patronymic[:-1] + 'у'
    elif patronymic.endswith('вна') or patronymic.endswith('ична'):
        patronymic_declined = patronymic[:-1] + 'у'
    else:
        patronymic_declined = patronymic + 'а'
    
    return f"{surname_declined} {name_declined} {patronymic_declined}"


def decline_court_name_genitive(court_name: str) -> str:
    """Склонение названия суда в родительный падеж"""
    
    # Арбитражный суд города Москвы -> Арбитражного суда города Москвы
    if 'Арбитражный суд города Москвы' in court_name:
        return 'Арбитражного суда города Москвы'
    
    # Арбитражный суд города Санкт-Петербурга и Ленинградской области -> Арбитражного суда города Санкт-Петербурга и Ленинградской области
    if 'Арбитражный суд города Санкт-Петербурга' in court_name:
        return 'Арбитражного суда города Санкт-Петербурга и Ленинградской области'
    
    # Универсальное правило для всех остальных:
    # Арбитражный суд [название] области/края/республики -> Арбитражного суда [название] области/края/республики
    if court_name.startswith('Арбитражный суд '):
        return court_name.replace('Арбитражный суд ', 'Арбитражного суда ')
    
    # Если не удалось определить - возвращаем как есть
    return court_name


def determine_court_by_address(address: str) -> Dict[str, str]:
    '''Определяет арбитражный суд по адресу регистрации'''
    
    address_lower = address.lower()
    
    courts_mapping = {
        'москва': {
            'name': 'Арбитражный суд города Москвы',
            'address': '115191, г. Москва, ул. Большая Тульская, д. 17'
        },
        'московская область': {
            'name': 'Арбитражный суд Московской области',
            'address': '107053, г. Москва, пр-т Академика Сахарова, д. 18'
        },
        'санкт-петербург': {
            'name': 'Арбитражный суд города Санкт-Петербурга и Ленинградской области',
            'address': '191015, г. Санкт-Петербург, Суворовский пр., д. 50-52'
        },
        'ленинградская область': {
            'name': 'Арбитражный суд города Санкт-Петербурга и Ленинградской области',
            'address': '191015, г. Санкт-Петербург, Суворовский пр., д. 50-52'
        },
        'новосибирск': {
            'name': 'Арбитражный суд Новосибирской области',
            'address': '630102, г. Новосибирск, ул. Нижегородская, д. 6'
        },
        'екатеринбург': {
            'name': 'Арбитражный суд Свердловской области',
            'address': '620075, г. Екатеринбург, ул. Шарташская, д. 4'
        },
        'казань': {
            'name': 'Арбитражный суд Республики Татарстан',
            'address': '420107, г. Казань, ул. Право-Булачная, д. 34/2'
        },
        'нижний новгород': {
            'name': 'Арбитражный суд Нижегородской области',
            'address': '603006, г. Нижний Новгород, Кремль, корп. 4'
        },
        'челябинск': {
            'name': 'Арбитражный суд Челябинской области',
            'address': '454000, г. Челябинск, ул. Воровского, д. 2'
        },
        'омск': {
            'name': 'Арбитражный суд Омской области',
            'address': '644024, г. Омск, ул. Учебная, д. 51'
        },
        'самара': {
            'name': 'Арбитражный суд Самарской области',
            'address': '443045, г. Самара, ул. Авроры, д. 148'
        },
        'ростов': {
            'name': 'Арбитражный суд Ростовской области',
            'address': '344002, г. Ростов-на-Дону, ул. Станиславского, д. 8'
        },
        'уфа': {
            'name': 'Арбитражный суд Республики Башкортостан',
            'address': '450057, г. Уфа, ул. Октябрьской революции, д. 63а'
        },
        'красноярск': {
            'name': 'Арбитражный суд Красноярского края',
            'address': '660049, г. Красноярск, пр. Мира, д. 63'
        },
        'пермь': {
            'name': 'Арбитражный суд Пермского края',
            'address': '614990, г. Пермь, ул. Екатерининская, д. 177'
        },
        'воронеж': {
            'name': 'Арбитражный суд Воронежской области',
            'address': '394030, г. Воронеж, ул. Свободы, д. 45'
        },
        'волгоград': {
            'name': 'Арбитражный суд Волгоградской области',
            'address': '400005, г. Волгоград, ул. Богунская, д. 12'
        },
        'краснодар': {
            'name': 'Арбитражный суд Краснодарского края',
            'address': '350063, г. Краснодар, ул. Постовая, д. 32'
        },
        'саратов': {
            'name': 'Арбитражный суд Саратовской области',
            'address': '410002, г. Саратов, ул. Бабушкин взвоз, д. 1'
        },
        'тюмень': {
            'name': 'Арбитражный суд Тюменской области',
            'address': '625000, г. Тюмень, ул. Первомайская, д. 20'
        },
        'тула': {
            'name': 'Арбитражный суд Тульской области',
            'address': '300041, г. Тула, ул. Староникитская, д. 1'
        },
        'иркутск': {
            'name': 'Арбитражный суд Иркутской области',
            'address': '664025, г. Иркутск, бул. Гагарина, д. 70'
        },
        'владивосток': {
            'name': 'Арбитражный суд Приморского края',
            'address': '690091, г. Владивосток, ул. Светланская, д. 54'
        },
        'хабаровск': {
            'name': 'Арбитражный суд Хабаровского края',
            'address': '680000, г. Хабаровск, ул. Ленина, д. 37'
        },
    }
    
    for region, court in courts_mapping.items():
        if region in address_lower:
            return court
    
    return {
        'name': 'Арбитражный суд (не удалось определить автоматически)',
        'address': 'Укажите адрес суда вручную'
    }


def generate_docx_document(
    personal: Dict[str, Any],
    credit: Dict[str, Any],
    income: Dict[str, Any],
    property: Dict[str, Any],
    additional: Dict[str, Any] = None,
    benefits: Dict[str, Any] = None,
    children: Dict[str, Any] = None,
    transactions: Dict[str, Any] = None,
    debt_reason: Dict[str, Any] = None,
    appendices: Dict[str, Any] = None
) -> str:
    '''Генерирует DOCX документ по шаблону и возвращает base64'''
    
    doc = Document()
    
    if additional is None:
        additional = {}
    
    passport = personal.get('passport', {})
    registration = personal.get('registration', {})
    creditors = credit.get('creditors', [])
    
    # Пересчитываем total_debt из реальных данных кредитов
    total_debt = 0
    for creditor in creditors:
        credits_list = creditor.get('credits', [])
        for credit_item in credits_list:
            total_debt += credit_item.get('debt', 0)
    
    registration_address = registration.get('address', '')
    auto_court = determine_court_by_address(registration_address)
    
    court_name = additional.get('courtName', auto_court['name'])
    court_address = additional.get('courtAddress', auto_court['address'])
    phone = personal.get('phone', 'Место для ввода текста.')
    email = personal.get('email', 'Место для ввода текста.')
    
    # Шапка документа с форматированием
    def add_header_paragraph(text):
        p = doc.add_paragraph(text)
        p_format = p.paragraph_format
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)
        p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_format.left_indent = Cm(7)
        return p
    
    add_header_paragraph(f"В {court_name}")
    add_header_paragraph(f"Адрес: {court_address}")
    add_header_paragraph("")
    
    add_header_paragraph(f"Заявитель (Должник):")
    add_header_paragraph(f"{personal.get('fullName', 'Место для ввода текста.')}")
    add_header_paragraph(f"Адрес: {registration.get('address', 'Место для ввода текста.')}")
    add_header_paragraph("")
    
    add_header_paragraph(f"Паспорт: серия {passport.get('series', 'Место для ввода текста.')} номер {passport.get('number', 'Место для ввода текста.')}")
    add_header_paragraph(f"выдан: {passport.get('issuedBy', 'Место для ввода текста.')}")
    add_header_paragraph(f"дата выдачи: {passport.get('issueDate', 'Место для ввода текста.')}")
    add_header_paragraph(f"код подразделения: {passport.get('code', 'Место для ввода текста.')}")
    add_header_paragraph(f"тел. {phone}")
    add_header_paragraph(f"e-mail: {email}")
    add_header_paragraph("")
    add_header_paragraph("")
    
    add_header_paragraph("Кредиторы:")
    for idx, creditor in enumerate(creditors, 1):
        add_header_paragraph(f"{idx}. {creditor.get('name', 'Место для ввода текста.')}")
        add_header_paragraph(f"ИНН {creditor.get('inn', 'Место для ввода текста.')}")
        legal_address = creditor.get('legalAddress', 'Место для ввода текста.')
        add_header_paragraph(f"Юридический адрес: {legal_address}")
        add_header_paragraph("")
    
    doc.add_paragraph()
    
    title = doc.add_heading("ЗАЯВЛЕНИЕ ГРАЖДАНИНА О ПРИЗНАНИИ БАНКРОТОМ", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Основной текст заявления с форматированием
    p_main = doc.add_paragraph(f"{personal.get('fullName', 'ФИО')} (далее – «Должник») обращается в суд с заявлением о признании его банкротом, поскольку имеются обязательства на сумму, превышающую 500 000 рублей и эти обязательства не исполнены Должником в течение трех месяцев с даты, когда они должны были быть исполнены. Кроме того, удовлетворение требований одного кредитора или нескольких кредиторов Должника приведет к невозможности исполнения Должником денежных обязательств в полном объеме перед другими кредиторами.")
    p_main_format = p_main.paragraph_format
    p_main_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_main_format.first_line_indent = Cm(1)
    p_main_format.space_before = Pt(0)
    p_main_format.space_after = Pt(0)
    p_main_format.line_spacing = 1.0
    
    p_debt = doc.add_paragraph(f"По состоянию на дату подачи заявления размер непогашенной задолженности Должника перед Кредиторами, составляет {format_number(total_debt)} рублей ({format_number(total_debt)} рублей 00 копеек) из которых:")
    p_debt_format = p_debt.paragraph_format
    p_debt_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_debt_format.first_line_indent = Cm(1)
    p_debt_format.space_before = Pt(0)
    p_debt_format.space_after = Pt(0)
    p_debt_format.line_spacing = 1.0
    
    for idx, creditor in enumerate(creditors, 1):
        doc.add_paragraph(f"{idx}. Перед {creditor.get('name', 'Место для ввода текста.')}:")
        credits_list = creditor.get('credits', [])
        if credits_list:
            for credit_item in credits_list:
                contract_num = credit_item.get('contractNumber', 'Место для ввода текста.')
                contract_date = credit_item.get('date', 'Место для ввода текста.')
                debt = credit_item.get('debt', 0)
                amount = credit_item.get('amount', 0)
                p_credit = doc.add_paragraph(f"- по Кредитному договору № {contract_num} от {contract_date} г. общая сумма задолженности составляет {format_number(debt)} рублей, в том числе сумма основного долга {format_number(amount)} рублей, сумма задолженности по начисленным процентам {format_number(debt - amount)} рублей.")
                p_credit_format = p_credit.paragraph_format
                p_credit_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_credit_format.first_line_indent = Cm(1)
                p_credit_format.space_before = Pt(0)
                p_credit_format.space_after = Pt(0)
                p_credit_format.line_spacing = 1.0
        else:
            p_credit_placeholder = doc.add_paragraph(f"- по Кредитному договору № Место для ввода текста. от Место для ввода текста. г. общая сумма задолженности составляет Место для ввода текста. рублей, в том числе сумма основного долга Место для ввода текста. рублей, сумма задолженности по начисленным процентам Место для ввода текста. рублей.")
            p_credit_placeholder_format = p_credit_placeholder.paragraph_format
            p_credit_placeholder_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_credit_placeholder_format.first_line_indent = Cm(1)
            p_credit_placeholder_format.space_before = Pt(0)
            p_credit_placeholder_format.space_after = Pt(0)
            p_credit_placeholder_format.line_spacing = 1.0
    
    monthly_income = income.get('monthlyIncome', 0) if income else 0
    yearly_income = income.get('lastYear', 0) if income else 0
    no_income = income.get('noIncome', False) if income else False
    
    if no_income or monthly_income == 0:
        p_no_income = doc.add_paragraph(f"Должник не имеет возможности удовлетворить требования Кредиторов в полном объеме, ввиду того, что объем всех ежемесячных обязательных платежей Должника в пользу Кредиторов (более {format_number(total_debt / 12)} рублей) превышает доход Должника. Должник не имеет постоянного источника дохода и является безработным.")
        p_no_income_format = p_no_income.paragraph_format
        p_no_income_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_no_income_format.first_line_indent = Cm(1)
        p_no_income_format.space_before = Pt(0)
        p_no_income_format.space_after = Pt(0)
        p_no_income_format.line_spacing = 1.0
        
        p = doc.add_paragraph("Должник не имеет дохода и является безработным, о чем свидетельствует справка из Центра занятости населения.")
        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_format.first_line_indent = Cm(1)
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)
        p_format.line_spacing = 1.0
    else:
        p1 = doc.add_paragraph(f"Должник не имеет возможности удовлетворить требования Кредиторов в полном объеме, ввиду того, что объем всех ежемесячных обязательных платежей Должника в пользу Кредиторов (более {format_number(total_debt / 12)} рублей) значительно превышает средний размер заработной платы Должника за последний год, что согласно справки 2 НДФЛ составляет {format_number(monthly_income)} рублей. ")
        p1_format = p1.paragraph_format
        p1_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p1_format.first_line_indent = Cm(1)
        p1_format.space_before = Pt(0)
        p1_format.space_after = Pt(0)
        p1_format.line_spacing = 1.0
        
        debt_reason_text = "Такая ситуация возникла в связи с тем, что Должник не рассчитал свои силы и возможности, кредиты приобретались при рождении детей и тратились на семейные нужды, впоследствии Должником была потеряна работа."
        if debt_reason and debt_reason.get('reason'):
            debt_reason_text = debt_reason.get('reason')
        
        p2 = doc.add_paragraph(debt_reason_text)
        p2_format = p2.paragraph_format
        p2_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p2_format.first_line_indent = Cm(1)
        p2_format.space_before = Pt(0)
        p2_format.space_after = Pt(0)
        p2_format.line_spacing = 1.0
    if benefits and benefits.get('isLowIncome'):
        cert_num = benefits.get('certificateNumber', 'Место для ввода текста.')
        cert_date = benefits.get('certificateDate', 'Место для ввода текста.')
        special_status = benefits.get('specialStatus', 'матерью одиночкой (инвалидом, пр.)')
        benefits_text = f"На основании справки № {cert_num} от {cert_date} г. Должник отнесен к категории малоимущих граждан. Кроме того, Должник является {special_status}."
        p3 = doc.add_paragraph(benefits_text)
        p3_format = p3.paragraph_format
        p3_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p3_format.first_line_indent = Cm(1)
        p3_format.space_before = Pt(0)
        p3_format.space_after = Pt(0)
        p3_format.line_spacing = 1.0
    
    if children and not children.get('noChildren', True):
        children_list = children.get('children', [])
        if children_list:
            for child in children_list:
                child_name = child.get('fullName', 'ФИО')
                birth_date = child.get('birthDate', 'года рождения')
                
                # Определяем пол по отчеству (последнее слово в ФИО)
                name_parts = child_name.split()
                is_male = True  # по умолчанию мужской род
                if len(name_parts) >= 3:
                    patronymic = name_parts[2].lower()
                    # Женские отчества заканчиваются на 'вна' или 'чна'
                    if patronymic.endswith('вна') or patronymic.endswith('ична'):
                        is_male = False
                
                # Склонение слов в зависимости от пола
                gender_word = 'несовершеннолетний' if is_male else 'несовершеннолетняя'
                living_with = 'проживающий вместе с Должником' if is_male else 'проживающая вместе с Должником'
                living_apart = 'проживающий отдельно от Должника' if is_male else 'проживающая отдельно от Должника'
                lives_with = living_with if child.get('livesWithDebtor', False) else living_apart
                
                p4 = doc.add_paragraph(f"На иждивении у Должника находится {gender_word} {child_name} {birth_date} года рождения, {lives_with}.")
                p4_format = p4.paragraph_format
                p4_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p4_format.first_line_indent = Cm(1)
                p4_format.space_before = Pt(0)
                p4_format.space_after = Pt(0)
                p4_format.line_spacing = 1.0
    
    if children and children.get('alimonyInfo', {}).get('isPayingAlimony', False):
        alimony = children.get('alimonyInfo', {})
        doc_type = alimony.get('documentType', 'notarial')
        
        if doc_type == 'notarial':
            notary_date = alimony.get('notaryDate', '01.01.2021')
            child_name = alimony.get('childFullName', 'ФИО')
            amount = format_number(alimony.get('monthlyAmount', 0))
            alimony_text = f"В соответствии с Соглашением об оплате алиментов на содержание несовершеннолетнего ребенка, удостоверенного нотариусом {notary_date} г., заключенного между Должником и матерью несовершеннолетнего ребенка Должника {child_name}, Должник обязался выплачивать в пользу своего несовершеннолетнего ребенка алименты в размере {amount} рублей ежемесячно до достижения ребенком восемнадцатилетнего возраста."
        elif doc_type == 'court':
            doc_details = alimony.get('documentDetails', 'решению суда')
            alimony_text = f"В соответствии с {doc_details} Должник обязан выплачивать алименты на содержание несовершеннолетнего ребенка."
        else:
            other_details = alimony.get('otherDetails', 'документу об алиментах')
            alimony_text = f"В соответствии с {other_details} Должник обязан выплачивать алименты на содержание несовершеннолетнего ребенка."
        
        p5 = doc.add_paragraph(alimony_text)
        p5_format = p5.paragraph_format
        p5_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p5_format.first_line_indent = Cm(1)
        p5_format.space_before = Pt(0)
        p5_format.space_after = Pt(0)
        p5_format.line_spacing = 1.0
    
    if property and property.get('realEstate'):
        real_estate = property.get('realEstate', [])
        if real_estate:
            # Обрабатываем все объекты недвижимости
            for idx, item in enumerate(real_estate):
                area_text = str(item.get('area', ''))
                land_area = item.get('landArea')
                
                # Формируем текст о земельном участке только если он указан
                land_text = ""
                if land_area and land_area != 0:
                    land_text = f" с земельным участком площадью {land_area} кв. м."
                
                if item.get('isSoleResidence', False):
                    p6 = doc.add_paragraph(f"В собственности Должника находится {item.get('type', 'недвижимость')}, общей площадью {area_text} кв. м.{land_text} по адресу: {item.get('address', 'Место для ввода текста.')}, являющийся единственным пригодным для постоянного проживания помещением для него и членов его семьи и не подлежит реализации.")
                else:
                    p6 = doc.add_paragraph(f"В собственности Должника находится {item.get('type', 'недвижимость')}, общей площадью {area_text} кв. м.{land_text} по адресу: {item.get('address', 'Место для ввода текста.')}.")
                
                p6_format = p6.paragraph_format
                p6_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p6_format.first_line_indent = Cm(1)
                p6_format.space_before = Pt(0)
                p6_format.space_after = Pt(0)
                p6_format.line_spacing = 1.0
    else:
        p6 = doc.add_paragraph("В собственности Должника находится недвижимое имущество, являющееся единственным пригодным для постоянного проживания помещением для него и членов его семьи.")
        p6_format = p6.paragraph_format
        p6_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p6_format.first_line_indent = Cm(1)
        p6_format.space_before = Pt(0)
        p6_format.space_after = Pt(0)
        p6_format.line_spacing = 1.0
    
    if property and property.get('vehicles'):
        vehicles = property.get('vehicles', [])
        if vehicles:
            # Обрабатываем все транспортные средства
            for idx, vehicle in enumerate(vehicles):
                vehicle_brand = vehicle.get('brand', 'Место для ввода текста.')
                vehicle_model = vehicle.get('model', 'Место для ввода текста.')
                vehicle_year = vehicle.get('year', 'Место для ввода текста.')
                vehicle_reg = vehicle.get('registrationNumber', 'Место для ввода текста.')
                vehicle_vin = vehicle.get('vin', 'Место для ввода текста.')
                
                intro = "Кроме того, в собственности Должника находится" if idx == 0 else "Также в собственности Должника находится"
                p7 = doc.add_paragraph(f"{intro} автомобиль марки {vehicle_brand} {vehicle_model} {vehicle_year} г., государственный регистрационный номер: {vehicle_reg}, идентификационный номер VIN: {vehicle_vin}.")
                p7_format = p7.paragraph_format
                p7_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p7_format.first_line_indent = Cm(1)
                p7_format.space_before = Pt(0)
                p7_format.space_after = Pt(0)
                p7_format.line_spacing = 1.0
    
    # Добавляем фразу об отсутствии имущества только если:
    # - стоит галочка "Имущество отсутствует" ИЛИ
    # - есть ровно один объект недвижимости с галочкой "единственное жилье" и нет транспортных средств
    real_estate = property.get('realEstate', []) if property else []
    vehicles = property.get('vehicles', []) if property else []
    no_property_checked = property.get('noProperty', False) if property else False
    
    should_show_no_property = (
        no_property_checked or
        (len(real_estate) == 1 and 
         real_estate[0].get('isSoleResidence', False) and 
         len(vehicles) == 0)
    )
    
    if should_show_no_property:
        p8 = doc.add_paragraph("Какое-либо имущество, на которое может быть обращено взыскание в соответствии с действующим законодательством, у Должника отсутствует.")
        p8_format = p8.paragraph_format
        p8_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p8_format.first_line_indent = Cm(1)
        p8_format.space_before = Pt(0)
        p8_format.space_after = Pt(0)
        p8_format.line_spacing = 1.0
    
    def add_body_paragraph(text):
        p = doc.add_paragraph(text)
        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_format.first_line_indent = Cm(1)
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)
        p_format.line_spacing = 1.0
        return p
    
    add_body_paragraph('В соответствии с п. 1 ст. 213.3 Федерального закона «О несостоятельности (банкротстве)» №127-ФЗ от 16.10.2002 (далее также – «Закон о банкротстве»), правом на обращение в суд с заявлением о признании гражданина банкротом обладают гражданин, конкурсный кредитор, уполномоченный орган.')
    add_body_paragraph('Согласно п. 1 с. 213.3 Закона о банкротстве, заявление о признании гражданина банкротом принимается судом при условии, что требования к гражданину составляют не менее чем пятьсот тысяч рублей и указанные требования не исполнены в течение трех месяцев с даты, когда они должны быть исполнены, если иное не предусмотрено Законом о банкротстве.')
    add_body_paragraph('В соответствии с п. 1 ст. 213.4 Закона о банкротстве, гражданин обязан обратиться в суд с заявлением о признании его банкротом в случае, если удовлетворение требований одного кредитора или нескольких кредиторов приводит к невозможности исполнения гражданином денежных обязательств и (или) обязанности по уплате обязательных платежей в полном объеме перед другими кредиторами и размер таких обязательств и обязанности в совокупности составляет не менее чем пятьсот тысяч рублей.')
    add_body_paragraph('Согласно п. 2 ст. 213.4 Закона о банкротстве гражданин вправе подать в арбитражный суд заявление о признании его банкротом в случае предвидения банкротства при наличии обстоятельств, очевидно свидетельствующих о том, что он не в состоянии исполнить денежные обязательства и (или) обязанность по уплате обязательных платежей в установленный срок, при этом гражданин отвечает признакам неплатежеспособности и (или) признакам недостаточности имущества.')
    add_body_paragraph('Согласно п.3 ст.213.6 Закона о банкротстве, под неплатежеспособностью гражданина понимается его неспособность удовлетворить в полном объеме требования кредиторов по денежным обязательствам и (или) исполнить обязанность по уплате обязательных платежей.')
    add_body_paragraph('Если не доказано иное, гражданин предполагается неплатежеспособным при условии, что имеет место хотя бы одно из следующих обстоятельств:')
    add_body_paragraph('- гражданин прекратил расчеты с кредиторами, то есть перестал исполнять денежные обязательства и (или) обязанность по уплате обязательных платежей, срок исполнения которых наступил;')
    add_body_paragraph('- более чем десять процентов совокупного размера денежных обязательств и (или) обязанности по уплате обязательных платежей, которые имеются у гражданина и срок исполнения которых наступил, не исполнены им в течение более чем одного месяца со дня, когда такие обязательства должны быть исполнены;')
    add_body_paragraph('- размер задолженности гражданина превышает стоимость его имущества, в том числе права требования;')
    add_body_paragraph('- наличие постановления об окончании исполнительного производства в связи с тем, что у гражданина отсутствует имущество, на которое может быть обращено взыскание.')
    
    add_body_paragraph("На момент подачи настоящего заявления:")
    add_body_paragraph('- существуют обстоятельства, очевидно свидетельствующие о том, что Должник не в состоянии исполнить денежные обязательства и (или) обязанность по уплате обязательных платежей в установленный срок;')
    add_body_paragraph(f"- сумма задолженности Должника перед Кредиторами составляет {format_number(total_debt)} рублей;")
    add_body_paragraph('- срок, в течение которого Должником не были исполнены обязательства, превышает 3 (три) месяца с момента наступления даты их исполнения.')
    add_body_paragraph('- размер задолженности Должника перед Кредитором превышает стоимость его имущества, на которое может быть обращено взыскание в соответствии с действующим законодательством (п. 3 ст. 213.25 Закона о банкротстве и ст. 446 ГПК РФ).')
    
    add_body_paragraph('Должник не привлекался к административной ответственности за мелкое хищение, умышленное уничтожение или повреждение имущества, неправомерные действия при банкротстве, фиктивное или преднамеренное банкротство.')
    add_body_paragraph('Кроме того, в настоящий момент отсутствуют сведения об известных Должнику уголовных и административных делах в отношении него, а также о наличии неснятой или непогашенной судимости.')
    
    if transactions and not transactions.get('noTransactions', True):
        transactions_list = transactions.get('transactions', [])
        if transactions_list:
            for idx, trans in enumerate(transactions_list, 1):
                trans_date = trans.get('date', '01.01.2021')
                trans_type = trans.get('type', 'купля-продажа')
                trans_desc = trans.get('description', 'сделка с имуществом')
                trans_amount = format_number(trans.get('amount', 0))
                trans_counterparty = trans.get('counterparty', 'Место для ввода текста.')
                p_trans = doc.add_paragraph(f'Должником в течение трех лет до даты подачи заявления была произведена сделка: {trans_type} от {trans_date} г. - {trans_desc}. Сумма сделки составила {trans_amount} рублей. Контрагент: {trans_counterparty}.')
                p_trans_format = p_trans.paragraph_format
                p_trans_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_trans_format.first_line_indent = Cm(1)
                p_trans_format.space_before = Pt(0)
                p_trans_format.space_after = Pt(0)
                p_trans_format.line_spacing = 1.0
            
            p10 = doc.add_paragraph('Иные сделки с недвижимым имуществом, ценными бумагами, долями в уставном капитале, транспортными средствами и сделок на сумму свыше трехсот тысяч рублей в течение трех лет до даты подачи настоящего заявления Должником не совершались.')
            p10_format = p10.paragraph_format
            p10_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p10_format.first_line_indent = Cm(1)
            p10_format.space_before = Pt(0)
            p10_format.space_after = Pt(0)
            p10_format.line_spacing = 1.0
    else:
        p10 = doc.add_paragraph('Сделки с недвижимым имуществом, ценными бумагами, долями в уставном капитале, транспортными средствами и сделок на сумму свыше трехсот тысяч рублей в течение трех лет до даты подачи настоящего заявления Должником не совершались.')
        p10_format = p10.paragraph_format
        p10_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p10_format.first_line_indent = Cm(1)
        p10_format.space_before = Pt(0)
        p10_format.space_after = Pt(0)
        p10_format.line_spacing = 1.0
    add_body_paragraph('В соответствии с общим смыслом п. 19 Постановления Пленума Верховного Суда РФ № 45 от 13.10.2015 г. «О некоторых вопросах, связанных с введением в действие процедур, применяемых в делах о несостоятельности (банкротстве) граждан» в качестве доказательства наличия у Должника денежных средств, достаточных для погашения расходов по делу о банкротстве, к настоящему заявлению приложен платежный документ об оплате в депозит Арбитражного суда денежных средств в размере 25 000 рублей.')
    add_body_paragraph('Таким образом, имеются признаки банкротства гражданина-должника, указанные в п. 3 ст. 213.6 Закона о банкротстве и основания для возбуждения судом дела о банкротстве в соответствии со статьями 213.3 и 213.4 Закона о банкротстве.')
    add_body_paragraph('На основании вышеизложенного, а также руководствуясь ст. ст. 6, 27, 38, 213.3, 213.4 Федерального закона «О несостоятельности (банкротстве)» от 26.10.2002 №127-ФЗ; ст. ст. 223, 224 АПК РФ,')
    
    doc.add_paragraph()
    prosh_heading = doc.add_paragraph("ПРОШУ:")
    prosh_heading.runs[0].bold = True
    prosh_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_body_paragraph(f"1) Признать {personal.get('fullName', 'Место для ввода текста.')} несостоятельным (банкротом) и ввести процедуру реализации имущества.")
    add_body_paragraph("2) Назначить финансового управляющего из числа членов следующих саморегулируемых организаций:")
    add_body_paragraph('• НПС СОПАУ "Альянс управляющих" (350015, Краснодарский край, г. Краснодар, ул.Северная, д.309);')
    add_body_paragraph('• Ассоциация арбитражных управляющих "ГАРАНТИЯ" (125167, г Москва, ул Викторенко, д.5, строение 1, эт. 2);')
    add_body_paragraph('• Ассоциация Арбитражных Управляющих "Содружество" (191124, г Санкт-Петербург, Санкт-Петербург, проспект Суворовский, д. 65, лит. Б, пом. 8-Н43);')
    add_body_paragraph('• АССОЦИАЦИЯ "МЕЖРЕГИОНАЛЬНАЯ САМОРЕГУЛИРУЕМАЯ ОРГАНИЗАЦИЯ ПРОФЕССИОНАЛЬНЫХ АРБИТРАЖНЫХ УПРАВЛЯЮЩИХ" (109240, г. Москва, Котельническая наб., д.17).')
    
    doc.add_paragraph()
    add_body_paragraph("Приложения:")
    
    # Формируем список приложений из переданных данных или используем дефолтный
    appendices_list = []
    
    if appendices and appendices.get('items'):
        items = appendices.get('items', [])
        # Фильтруем только включенные приложения
        included_items = [item for item in items if item.get('isIncluded', False)]
        
        for idx, item in enumerate(included_items, 1):
            title = item.get('title', 'Приложение')
            pages = item.get('pages', 'Место для ввода текста.')
            
            # Форматируем строку приложения
            if pages and pages.strip():
                # Добавляем склонение слова "лист"
                pages_formatted = get_page_word(pages)
                appendices_list.append(f"{idx}. {title} {pages_formatted};")
            else:
                appendices_list.append(f"{idx}. {title} Место для ввода текста.;")
        
        # Добавляем выписки из ЕГРЮЛ для кредиторов
        start_idx = len(included_items) + 1
        for creditor_idx in range(len(creditors)):
            appendices_list.append(f"{start_idx + creditor_idx}. Выписка из ЕГРЮЛ на кредитора {creditor_idx + 1};")
    else:
        # Дефолтный список если данные не переданы
        appendices_list = [
            "1. Почтовые квитанции о направлении копии настоящего заявления кредиторам, Место для ввода текста. листов.",
            "2. Документ, подтверждающий уплату государственной пошлины Место для ввода текста. лист;",
            "3. Документ, подтверждающий внесение денежных средств на выплату вознаграждения финансовому управляющему в депозит арбитражного суда Место для ввода текста. лист;",
            "4. Копия паспорта (всех страниц) Место для ввода текста. листов;",
            "5. Копия ИНН (свидетельства о постановке на учет в налоговом органе) (при наличии) Место для ввода текста. лист;",
            "6. Копия СНИЛС Место для ввода текста. лист;",
            "7. Копии справок, выданных кредиторами о задолженности на Место для ввода текста. листах;",
            "8. Копии кредитных договоров Место для ввода текста. штуки;",
            "9. Документы, подтверждающие наличие или отсутствие у гражданина статуса индивидуального предпринимателя, полученные не ранее чем за пять рабочих дней до даты подачи в суд заявления на Место для ввода текста. листе;",
            "10. Список кредиторов и должников гражданина по форме Приложения № 1 к Приказу Минэкономразвития России от 05.08.2015 № 530;",
            "11. Опись имущества гражданина по форме Приложения № 2 к Приказу Минэкономразвития России от 05.08.2015 № 530;",
            "12. Справки 2 НДФЛ за трехлетний период, предшествующий дате подачи заявления о признании гражданина банкротом на Место для ввода текста. листах;",
            "13. Копии документов, подтверждающих право собственности гражданина на имущество на Место для ввода текста. листах (при наличии);",
            "14. Копии документов о совершавшихся гражданином в течение трех лет до даты подачи заявления сделках с недвижимым имуществом, транспортными средствами и сделках на сумму свыше трехсот тысяч рублей (при наличии);",
            "15. Копия свидетельства пенсионного страхования (при наличии);",
            "16. Копия решения о признании гражданина безработным, выданная государственной службой занятости населения (при наличии);",
            "17. Копия свидетельства о заключении брака (при наличии заключенного и не расторгнутого на дату подачи заявления брака);",
            "18. Копия свидетельства о расторжении брака, если оно выдано в течение трех лет до даты подачи заявления (при наличии);",
            "19. Копия брачного договора (при наличии);",
            "20. Копия соглашения или судебного акта о разделе общего имущества супругов, соответственно заключенного и принятого в течение трех лет до даты подачи заявления (при наличии);",
            "21. Копия свидетельства о рождении ребенка, если гражданин является его родителем, усыновителем или опекуном.",
        ]
        for creditor_idx in range(len(creditors)):
            appendices_list.append(f"{22 + creditor_idx}. Выписка из ЕГРЮЛ на кредитора {creditor_idx + 1};")
    
    for app in appendices_list:
        add_body_paragraph(app)
    
    doc.add_paragraph()
    add_body_paragraph(f"Должник: {personal.get('fullName', 'Место для ввода текста.')}")
    add_body_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return base64.b64encode(buffer.read()).decode('utf-8')


def generate_pdf_document(
    personal: Dict[str, Any],
    credit: Dict[str, Any],
    income: Dict[str, Any],
    property: Dict[str, Any],
    additional: Dict[str, Any] = None,
    benefits: Dict[str, Any] = None,
    children: Dict[str, Any] = None,
    transactions: Dict[str, Any] = None,
    debt_reason: Dict[str, Any] = None
) -> str:
    '''Генерирует PDF документ по шаблону и возвращает base64'''
    
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    
    y = height - 2*cm
    
    c.setFont("Helvetica", 10)
    
    if additional is None:
        additional = {}
    
    passport = personal.get('passport', {})
    registration = personal.get('registration', {})
    creditors = credit.get('creditors', [])
    
    # Пересчитываем total_debt из реальных данных кредитов
    total_debt = 0
    for creditor in creditors:
        credits_list = creditor.get('credits', [])
        for credit_item in credits_list:
            total_debt += credit_item.get('debt', 0)
    
    registration_address = registration.get('address', '')
    auto_court = determine_court_by_address(registration_address)
    
    court_name = additional.get('courtName', auto_court['name'])
    court_address = additional.get('courtAddress', auto_court['address'])
    phone = personal.get('phone', 'Место для ввода текста.')
    email = personal.get('email', 'Место для ввода текста.')
    
    c.drawString(2*cm, y, f"В {court_name}")
    y -= 0.5*cm
    c.drawString(2*cm, y, f"Адрес: {court_address}")
    y -= 1*cm
    
    c.drawString(2*cm, y, f"Заявитель (Должник):")
    y -= 0.5*cm
    c.drawString(2*cm, y, f"{personal.get('fullName', 'Место для ввода текста.')}")
    y -= 0.5*cm
    c.drawString(2*cm, y, f"Адрес: {registration.get('address', 'Место для ввода текста.')}")
    y -= 1*cm
    
    c.drawString(2*cm, y, f"Паспорт: серия {passport.get('series', 'Место для ввода текста.')} номер {passport.get('number', 'Место для ввода текста.')}")
    y -= 0.5*cm
    c.drawString(2*cm, y, f"выдан: {passport.get('issuedBy', 'Место для ввода текста.')}")
    y -= 0.5*cm
    c.drawString(2*cm, y, f"дата выдачи: {passport.get('issueDate', 'Место для ввода текста.')}")
    y -= 0.5*cm
    c.drawString(2*cm, y, f"код подразделения: {passport.get('code', 'Место для ввода текста.')}")
    y -= 1*cm
    
    c.drawString(2*cm, y, "Кредиторы:")
    y -= 0.5*cm
    
    for idx, creditor in enumerate(creditors, 1):
        c.drawString(2*cm, y, f"{idx}. {creditor.get('name', 'Место для ввода текста.')}")
        y -= 0.5*cm
        c.drawString(2*cm, y, f"ИНН {creditor.get('inn', 'Место для ввода текста.')}")
        y -= 0.5*cm
        legal_address = creditor.get('legalAddress', 'Место для ввода текста.')
        c.drawString(2*cm, y, f"Юридический адрес: {legal_address}")
        y -= 0.7*cm
        
        if y < 5*cm:
            c.showPage()
            y = height - 2*cm
            c.setFont("Helvetica", 10)
    
    y -= 1*cm
    c.setFont("Helvetica-Bold", 12)
    c.drawCentredString(width/2, y, "ЗАЯВЛЕНИЕ ГРАЖДАНИНА О ПРИЗНАНИИ БАНКРОТОМ")
    y -= 1.5*cm
    
    c.setFont("Helvetica", 10)
    
    text_lines = [
        f"{personal.get('fullName', 'ФИО')} (далее – «Должник») обращается в суд с заявлением",
        "о признании его банкротом, поскольку имеются обязательства на сумму,",
        "превышающую 500 000 рублей...",
        "",
        f"Общая задолженность: {format_number(total_debt)} рублей",
    ]
    
    for line in text_lines:
        if y < 3*cm:
            c.showPage()
            y = height - 2*cm
            c.setFont("Helvetica", 10)
        c.drawString(2*cm, y, line)
        y -= 0.5*cm
    
    y -= 1*cm
    c.setFont("Helvetica-Bold", 10)
    c.drawString(2*cm, y, "ПРОШУ:")
    y -= 0.7*cm
    
    c.setFont("Helvetica", 10)
    c.drawString(2*cm, y, f"1) Признать {personal.get('fullName', 'Место для ввода текста.')} несостоятельным (банкротом)")
    y -= 0.5*cm
    c.drawString(2*cm, y, "2) Назначить финансового управляющего")
    y -= 1.5*cm
    
    c.drawString(2*cm, y, f"Должник: {personal.get('fullName', 'Место для ввода текста.')}")
    y -= 0.7*cm
    c.drawString(2*cm, y, f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
    
    c.save()
    buffer.seek(0)
    return base64.b64encode(buffer.read()).decode('utf-8')


def generate_creditors_list_document(
    personal: Dict[str, Any],
    credit: Dict[str, Any]
) -> str:
    '''
    Генерирует Приложение №1 - Список кредиторов и должников гражданина
    по форме из Приказа Минэкономразвития от 05.08.2015 N 530
    '''
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.table import _Cell
    
    doc = Document()
    
    # Заголовок документа
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p_header.add_run("Приложение № 1\nк приказу Минэкономразвития России\nот 5 августа 2015 г. № 530")
    run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Заголовок
    title = doc.add_heading("Список кредиторов и должников гражданина", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(14)
        run.font.bold = True
    
    doc.add_paragraph()
    
    # Таблица с информацией о гражданине
    info_table = doc.add_table(rows=14, cols=2)
    info_table.style = 'Table Grid'
    
    # Заголовок таблицы
    info_header = info_table.rows[0].cells
    info_header[0].merge(info_header[1])
    info_header[0].text = "Информация о гражданине"
    info_header[0].paragraphs[0].runs[0].font.bold = True
    info_header[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Извлекаем данные
    full_name = personal.get('fullName', 'Ф.И.О.')
    full_name_parts = full_name.split()
    surname = full_name_parts[0] if len(full_name_parts) > 0 else 'обязательно'
    name = full_name_parts[1] if len(full_name_parts) > 1 else 'обязательно'
    patronymic = full_name_parts[2] if len(full_name_parts) > 2 else 'при наличии'
    
    birth_date = personal.get('birthDate', 'обязательно')
    birth_place = personal.get('birthPlace', 'обязательно')
    snils = personal.get('snils', 'обязательно')
    inn = personal.get('inn', 'при наличии')
    
    passport = personal.get('passport', {})
    passport_series = passport.get('series', 'обязательно')
    passport_number = passport.get('number', 'обязательно')
    
    registration = personal.get('registration', {})
    address_parts = registration.get('address', '').split(',')
    
    # Заполняем строки таблицы
    rows_data = [
        ('фамилия', surname),
        ('имя', name),
        ('отчество', patronymic),
        ('в случае изменения фамилии, имени, отчества указать прежние фамилии, имена, отчества', 'обязательно'),
        ('дата рождения', birth_date),
        ('место рождения', birth_place),
        ('СНИЛС', snils),
        ('ИНН', inn),
        ('документ, удостоверяющий личность', ''),
        ('вид документа', 'паспорт гражданина РФ'),
        ('серия (при наличии) и номер', f'{passport_series} {passport_number}'),
        ('адрес регистрации по месту жительства в Российской Федерации*', ''),
        ('субъект Российской Федерации', address_parts[0].strip() if len(address_parts) > 0 else 'обязательно'),
    ]
    
    for idx, (label, value) in enumerate(rows_data, 1):
        row = info_table.rows[idx].cells
        row[0].text = label
        row[1].text = value
        row[0].paragraphs[0].runs[0].font.size = Pt(9)
        if row[1].text:
            row[1].paragraphs[0].runs[0].font.size = Pt(9)
    
    doc.add_paragraph()
    
    # Функция для объединения ячеек
    def set_cell_border(cell, **kwargs):
        """Устанавливает границы ячейки"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            if edge in kwargs:
                edge_el = OxmlElement(f'w:{edge}')
                edge_el.set(qn('w:val'), 'single')
                edge_el.set(qn('w:sz'), '4')
                edge_el.set(qn('w:space'), '0')
                edge_el.set(qn('w:color'), '000000')
                tcBorders.append(edge_el)
        tcPr.append(tcBorders)
    
    # Раздел II: Сведения о кредиторах
    section_heading = doc.add_paragraph()
    run_heading = section_heading.add_run("II. Сведения о ")
    run_heading.bold = True
    run_creditors = section_heading.add_run("кредиторах")
    run_creditors.bold = True
    run_creditors.font.color.rgb = RGBColor(255, 0, 0)
    run_heading2 = section_heading.add_run(" гражданина")
    run_heading2.bold = True
    
    p_desc = doc.add_paragraph()
    run1 = p_desc.add_run("(по денежным обязательствам и (или) обязанности по уплате обязательных платежей, ")
    run2 = p_desc.add_run("которые возникли в результате осуществления гражданином предпринимательской деятельности")
    run2.font.color.rgb = RGBColor(255, 0, 0)
    run3 = p_desc.add_run(")")
    
    doc.add_paragraph()
    
    # Таблица 1: Денежные обязательства
    p_money = doc.add_paragraph()
    p_money.add_run("1").bold = True
    p_money_label = p_money.add_run("\t\tДенежные обязательства")
    
    # Таблица кредиторов с расширенными заголовками
    table = doc.add_table(rows=2, cols=7)
    table.style = 'Table Grid'
    
    # Первая строка заголовков (объединенные ячейки)
    header_row1 = table.rows[0].cells
    header_row1[0].text = "№ п/п"
    header_row1[1].text = "Содержание обязательства⁴"
    header_row1[2].text = "Кредитор⁵"
    header_row1[3].text = "Место нахождения (место жительства) кредитора"
    header_row1[4].text = "Основание возникновения⁶"
    header_row1[5].text = "Сумма обязательства"
    header_row1[6].text = "Штрафы, пени и иные санкции"
    
    # Вторая строка заголовков (подзаголовки для суммы)
    header_row2 = table.rows[1].cells
    header_row2[0].text = ""
    header_row2[1].text = ""
    header_row2[2].text = ""
    header_row2[3].text = "⁰⁰"
    header_row2[4].text = ""
    header_row2[5].text = "всего⁷"
    header_row2[6].text = ""
    
    # Объединяем ячейки для колонок без подзаголовков
    for idx in [0, 1, 2, 4, 6]:
        header_row1[idx].merge(header_row2[idx])
    
    # Форматирование заголовков
    for row in [header_row1, header_row2]:
        for cell in row:
            if cell.text:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Строка с пустыми значениями (1.1)
    empty_row = table.add_row().cells
    empty_row[0].text = "1.1"
    empty_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    empty_row[1].text = "0"
    empty_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    empty_row[2].text = "0"
    empty_row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    empty_row[3].text = "0"
    empty_row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    empty_row[4].text = "00"
    empty_row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    empty_row[5].text = "0"
    empty_row[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    empty_row[6].text = "0"
    empty_row[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    for cell in empty_row:
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    
    # Заполнение данными о кредиторах
    creditors = credit.get('creditors', [])
    row_num = 1
    total_debt = 0
    
    for creditor in creditors:
        creditor_name = creditor.get('name', '')
        creditor_inn = creditor.get('inn', '')
        creditor_address = creditor.get('legalAddress', 'не указано')
        
        credits_list = creditor.get('credits', [])
        for credit_item in credits_list:
            row = table.add_row().cells
            
            # № п/п
            row[0].text = f"1.{row_num}"
            row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Содержание обязательства
            row[1].text = "Кредитный договор"
            
            # Кредитор
            row[2].text = f"{creditor_name}\nИНН: {creditor_inn}"
            
            # Место нахождения
            row[3].text = creditor_address
            
            # Основание возникновения
            contract_num = credit_item.get('contractNumber', '')
            contract_date = credit_item.get('date', '')
            row[4].text = f"Договор № {contract_num}\nот {contract_date}"
            
            # Сумма обязательства всего
            debt = credit_item.get('debt', 0)
            row[5].text = format_number(debt)
            row[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            total_debt += debt
            
            # Штрафы
            row[6].text = "0"
            row[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Форматирование ячеек
            for cell in row:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
            
            row_num += 1
    
    doc.add_paragraph()
    
    # Раздел с обязательными платежами
    p_tax = doc.add_paragraph()
    p_tax.add_run("2\t\tОбязательные платежи").bold = True
    
    table_tax = doc.add_table(rows=2, cols=3)
    table_tax.style = 'Table Grid'
    
    headers_tax = table_tax.rows[0].cells
    headers_tax[0].text = "№ п/п"
    headers_tax[1].text = "Наименование налога, сбора или иного обязательного платежа"
    headers_tax[2].text = "Штрафы, пени и иные санкции"
    
    for cell in headers_tax:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Пустая строка (нет налоговых задолженностей)
    row_tax = table_tax.rows[1].cells
    row_tax[0].text = "2.1 | 0"
    row_tax[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row_tax[1].text = ""
    row_tax[2].text = "0"
    row_tax[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # Текст о предпринимательской деятельности
    p_entrepreneur = doc.add_paragraph()
    p_entrepreneur.add_run("Сведения о неисполненных обязательствах гражданина, которые возникли в результате осуществления гражданином предпринимательской деятельности (в том числе о передаче имущества в собственность, выполнении работ и оказании услуг и так далее):")
    p_entrepreneur.paragraph_format.first_line_indent = Cm(1)
    p_entrepreneur.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Линии для заполнения
    for _ in range(3):
        p_line = doc.add_paragraph("_" * 100)
        p_line.paragraph_format.space_before = Pt(0)
        p_line.paragraph_format.space_after = Pt(0)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Сноски внизу
    p_footnotes = doc.add_paragraph()
    p_footnotes.add_run("_" * 50)
    p_footnotes.paragraph_format.space_before = Pt(12)
    
    footnotes_text = [
        "* Указывается субъект обязательства (кредитор, заем, кредит).",
        "⁴ Указывается сумма срочных обязательств, кредитора, должника от остатка (основание — при наличии) для физического лица или наименование юридического лица.",
        "⁵ Указывается ИНН, адрес и регистрационный номер кредитора, если кредитор — физическое лицо или наименование юридического лица.",
        "⁶ Указывается суждение обязательства (например, заем, кредит).",
        "⁷ Указывается общая сумма основного обязательства, в том числе наложения платежи проценты. Для обязательства, выраженного в иностранной валюте, сумма указывается в рублях по курсу Банка России на дату составления списка кредиторов и должников гражданина (за исключением внутреннего картина), пени, проценты за просрочку платежа, учетные в сумме основной задолженности в виде финансовых санкций, начисленных на сумму основного обязательства). Для обязательства, выраженных в иностранной валюте, сумма указывается в рублях по курсу Банка России на дату составления списка кредиторов и должников гражданина.",
        "⁰⁰ Указывается сумма суммы обязательства (кредитор, срок, кредит)."
    ]
    
    for footnote in footnotes_text:
        p_fn = doc.add_paragraph(footnote)
        p_fn.paragraph_format.first_line_indent = Cm(0)
        p_fn.runs[0].font.size = Pt(8)
        p_fn.paragraph_format.space_before = Pt(2)
        p_fn.paragraph_format.space_after = Pt(2)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Подпись
    p_signature = doc.add_paragraph()
    p_signature.add_run(f"Гражданин: {full_name}")
    p_signature.add_run("\t\t_______________")
    
    p_date = doc.add_paragraph()
    p_date.add_run(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
    
    # Сохранение в base64
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return base64.b64encode(buffer.read()).decode('utf-8')


def generate_property_list_document(
    personal: Dict[str, Any],
    property: Dict[str, Any]
) -> str:
    '''
    Генерирует Приложение №2 - Опись имущества гражданина
    по форме из Приказа Минэкономразвития от 05.08.2015 N 530
    '''
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    doc = Document()
    
    # Заголовок документа
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p_header.add_run("Приложение № 2\nк приказу Минэкономразвития России\nот 5 августа 2015 г. № 530")
    run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Заголовок
    title = doc.add_heading("Опись имущества гражданина", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(14)
        run.font.bold = True
    
    doc.add_paragraph()
    
    # Таблица с информацией о гражданине
    info_table = doc.add_table(rows=14, cols=2)
    info_table.style = 'Table Grid'
    
    # Заголовок таблицы
    info_header = info_table.rows[0].cells
    info_header[0].merge(info_header[1])
    info_header[0].text = "Информация о гражданине"
    info_header[0].paragraphs[0].runs[0].font.bold = True
    info_header[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Извлекаем данные
    full_name = personal.get('fullName', '_________________________')
    full_name_parts = full_name.split()
    surname = full_name_parts[0] if len(full_name_parts) > 0 else 'обязательно'
    name = full_name_parts[1] if len(full_name_parts) > 1 else 'обязательно'
    patronymic = full_name_parts[2] if len(full_name_parts) > 2 else 'при наличии'
    
    birth_date = personal.get('birthDate', 'обязательно')
    birth_place = personal.get('birthPlace', 'обязательно')
    snils = personal.get('snils', 'обязательно')
    inn = personal.get('inn', 'при наличии')
    
    passport = personal.get('passport', {})
    passport_series = passport.get('series', 'обязательно')
    passport_number = passport.get('number', 'обязательно')
    
    registration = personal.get('registration', {})
    address = registration.get('address', 'обязательно')
    
    # Извлекаем субъект РФ из адреса (первая часть до запятой)
    address_parts = address.split(',')
    subject_rf = address_parts[0].strip() if len(address_parts) > 0 else 'обязательно'
    
    # Заполняем строки таблицы
    rows_data = [
        ('фамилия', surname),
        ('имя', name),
        ('отчество', patronymic),
        ('в случае изменения фамилии, имени, отчества указать прежние фамилии, имена, отчества', 'обязательно'),
        ('дата рождения', birth_date),
        ('место рождения', birth_place),
        ('СНИЛС', snils),
        ('ИНН', inn),
        ('документ, удостоверяющий личность', ''),
        ('вид документа', 'паспорт'),
        ('серия (при наличии) и номер', f'{passport_series} {passport_number}'),
        ('адрес регистрации по месту жительства в Российской Федерации*', ''),
        ('субъект Российской Федерации', subject_rf),
    ]
    
    for idx, (label, value) in enumerate(rows_data, 1):
        row = info_table.rows[idx].cells
        row[0].text = label
        row[1].text = value
        row[0].paragraphs[0].runs[0].font.size = Pt(9)
        if row[1].text:
            row[1].paragraphs[0].runs[0].font.size = Pt(9)
    
    doc.add_paragraph()
    
    # I. Недвижимое имущество
    section1_heading = doc.add_paragraph()
    section1_heading.add_run("I. Недвижимое имущество").bold = True
    
    table1 = doc.add_table(rows=2, cols=6)
    table1.style = 'Table Grid'
    
    # Первая строка заголовков
    headers1_row1 = table1.rows[0].cells
    headers1_row1[0].text = "№ п/п"
    headers1_row1[1].text = "Вид и наименование имущества"
    headers1_row1[2].text = "Вид собственности¹"
    headers1_row1[3].text = "Местонахождение (адрес)"
    headers1_row1[4].text = "Площадь (кв. м)"
    headers1_row1[5].text = "Сведения о залоге и залогодержателе ¹"
    
    # Вторая строка (подзаголовки) - оставляем пустой для большинства колонок
    headers1_row2 = table1.rows[1].cells
    
    # Объединяем ячейки для колонок без подзаголовков
    for idx in [0, 1, 2, 3, 4, 5]:
        headers1_row1[idx].merge(headers1_row2[idx])
    
    # Форматирование заголовков
    for cell in headers1_row1:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Добавляем данные о недвижимости (или пустую строку)
    real_estate = property.get('realEstate', []) if property else []
    if not real_estate:
        # Добавляем пустую строку если нет данных
        empty_row = table1.add_row().cells
        empty_row[0].text = ""
        empty_row[1].text = ""
        empty_row[2].text = ""
        empty_row[3].text = ""
        empty_row[4].text = ""
        empty_row[5].text = ""
    
    doc.add_paragraph()
    
    # II. Движимое имущество
    section2_heading = doc.add_paragraph()
    section2_heading.add_run("II. Движимое имущество").bold = True
    
    table2 = doc.add_table(rows=2, cols=5)
    table2.style = 'Table Grid'
    
    # Заголовки таблицы
    headers2_row1 = table2.rows[0].cells
    headers2_row1[0].text = "№ п/п"
    headers2_row1[1].text = "Вид, марка, модель транспортного средства, год изготовления"
    headers2_row1[2].text = "Идентификационный номер⁴"
    headers2_row1[3].text = "Вид собственности⁵"
    headers2_row1[4].text = "Сведения о залоге и залогодержателе⁸"
    
    headers2_row2 = table2.rows[1].cells
    # Объединяем ячейки
    for idx in [0, 1, 2, 3, 4]:
        headers2_row1[idx].merge(headers2_row2[idx])
    
    for cell in headers2_row1:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Пустая строка
    empty_row2 = table2.add_row().cells
    for cell in empty_row2:
        cell.text = ""
    
    doc.add_paragraph()
    
    # III. Сведения о счетах в банках и иных кредитных организациях
    section3_heading = doc.add_paragraph()
    section3_heading.add_run("III. Сведения о счетах в банках и иных кредитных организациях").bold = True
    
    table3 = doc.add_table(rows=2, cols=4)
    table3.style = 'Table Grid'
    
    headers3_row1 = table3.rows[0].cells
    headers3_row1[0].text = "№ п/п"
    headers3_row1[1].text = "Наименование и адрес банка или иной кредитной организации"
    headers3_row1[2].text = "Вид и валюта счета⁹"
    headers3_row1[3].text = "Дата открытия счета"
    
    headers3_row2 = table3.rows[1].cells
    headers3_row2[0].text = ""
    headers3_row2[1].text = ""
    headers3_row2[2].text = ""
    headers3_row2[3].text = "-"
    
    for idx in [0, 1, 2]:
        headers3_row1[idx].merge(headers3_row2[idx])
    
    for cell in headers3_row1:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # IV. Акции и иное участие в коммерческих организациях
    section4_heading = doc.add_paragraph()
    section4_heading.add_run("IV. Акции и иное участие в коммерческих организациях").bold = True
    doc.add_paragraph()
    
    # V. Иные ценные бумаги
    section5_heading = doc.add_paragraph()
    section5_heading.add_run("V. Иные ценные бумаги").bold = True
    
    table5 = doc.add_table(rows=2, cols=5)
    table5.style = 'Table Grid'
    
    headers5_row1 = table5.rows[0].cells
    headers5_row1[0].text = "№ п/п"
    headers5_row1[1].text = "Вид ценной бумаги¹⁵"
    headers5_row1[2].text = "Лицо, выпустившее ценную бумагу"
    headers5_row1[3].text = "Номинальная величина обязательства (руб.)"
    headers5_row1[4].text = "Общее количество"
    
    headers5_row2 = table5.rows[1].cells
    headers5_row2[0].text = "5.1"
    headers5_row2[1].text = ""
    headers5_row2[2].text = ""
    headers5_row2[3].text = ""
    headers5_row2[4].text = ""
    
    for idx in [0, 1, 2, 3, 4]:
        headers5_row1[idx].merge(headers5_row2[idx])
    
    for cell in headers5_row1:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Строки 5.2, 5.3
    for num in ['5.2', '5.3']:
        row = table5.add_row().cells
        row[0].text = num
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # VI. Сведения о наличных денежных средствах и ином ценном имуществе
    section6_heading = doc.add_paragraph()
    section6_heading.add_run("VI. Сведения о наличных денежных средствах и ином ценном имуществе").bold = True
    
    table6 = doc.add_table(rows=2, cols=4)
    table6.style = 'Table Grid'
    
    headers6 = table6.rows[0].cells
    headers6[0].text = "№ п/п"
    headers6[1].text = "Вид и наименование имущества"
    headers6[2].text = "Стоимость (сумма в валюте)¹⁶"
    headers6[3].text = "Место нахождения/ место хранения¹⁸ (адрес)"
    
    headers6_row2 = table6.rows[1].cells
    for idx in range(4):
        headers6[idx].merge(headers6_row2[idx])
    
    for cell in headers6:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Подтверждение достоверности
    p_confirm = doc.add_paragraph()
    p_confirm.add_run("Достоверность и полноту настоящих сведений подтверждаю.")
    p_confirm.paragraph_format.first_line_indent = Cm(1)
    
    doc.add_paragraph()
    
    # Подпись и дата
    p_date_line = doc.add_paragraph()
    p_date_line.add_run(f'"___" _________ 20___ г.')
    p_date_line.add_run("\t\t\t")
    p_date_line.add_run("_(подпись гражданина)_")
    p_date_line.add_run("\t\t")
    p_date_line.add_run("_(расшифровка подписи)_")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Сноски
    p_footnotes_separator = doc.add_paragraph()
    p_footnotes_separator.add_run("_" * 50)
    
    footnotes = [
        "¹ Указывается вид собственности (индивидуальная, долевая, общая); для совместной собственности указывается также иные лица (фамилия, имя и отчество (последнее - при наличии) или наименование), в собственности которых находится имущество; для долевой собственности указывается размер доли гражданина.",
        "⁴ Указывается при наличии документов, содержащих сведения о стоимости имущества (например, отчет о стоимости имущества, подготовленный независимой оценочной компанией либо данные оценки для целей налогообложения).",
        "⁵ Указывается лицо, в пользу которого установлен залог, и сведения о договоре залога (номер и дата).",
        "⁸ Указывается полное наименование юридического лица или фамилия, имя, отчество (последнее - при наличии), участником (акционером) которого является гражданин, а также размер доли участия в голосующих акциях или уставном капитале юридического лица либо количество принадлежащих акций (для акционера).",
        "⁹ Указывается сведения о депозитах, включая номинальные (опосредованно, договор, депозит, счет на котором возложит исполнение иска или иную вещь на лицо - владельца либо пользователя, либо и изложения акций, указываемых в разделе IV (настоящий пункт для выгодоприобретателя).",
        "¹⁵ Указывается вид ценной бумаги, например, вексель, чек, облигация, опцион, фьючерсный договор, коносамент, складское свидетельство, закладная, инвестиционный пай, сертификат участия, сертификат долей, депозит, иной депозит.",
        "¹⁶ Указывается сумма в валюте, включенная при внесении документов, содержащих сведения о стоимости имущества (например, отчет о стоимости имущества, подготовленный независимой оценочной компанией либо данные кассового аппарата, которые подтверждает акт или иное принадлежности выдаче в курсе Банка России на дату составления описи имущества гражданина.",
        "¹⁸ Указывается сведения о договоре аренды, лизинга, прочей сделке, на основании которой возложит право или же лицо обязано передачу в иностранном банке - на рабочее указывается также или факт банка и место курса Банка России на дату составления описи имущества гражданина."
    ]
    
    for footnote in footnotes:
        p_fn = doc.add_paragraph(footnote)
        p_fn.runs[0].font.size = Pt(8)
        p_fn.paragraph_format.space_before = Pt(2)
        p_fn.paragraph_format.space_after = Pt(2)
    
    # Сохранение в base64
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return base64.b64encode(buffer.read()).decode('utf-8')


def generate_bankruptcy_application(
    personal: Dict[str, Any],
    credit: Dict[str, Any],
    income: Dict[str, Any],
    property: Dict[str, Any]
) -> str:
    '''Генерирует текст заявления о банкротстве (устарело, используйте DOCX/PDF)'''
    return "Используйте формат DOCX или PDF для генерации документа"


def generate_attachment_motion_document(
    personal: Dict[str, Any],
    additional: Dict[str, Any],
    motion_data: Dict[str, Any]
) -> str:
    '''Генерирует DOCX документ ходатайства о приобщении документов'''
    
    doc = Document()
    
    if additional is None:
        additional = {}
    
    passport = personal.get('passport', {})
    registration = personal.get('registration', {})
    
    registration_address = registration.get('address', '')
    auto_court = determine_court_by_address(registration_address)
    
    court_name = additional.get('courtName', auto_court['name'])
    court_address = additional.get('courtAddress', auto_court['address'])
    phone = personal.get('phone', 'Место для ввода текста.')
    email = personal.get('email', 'Место для ввода текста.')
    full_name = personal.get('fullName', 'Место для ввода текста.')
    
    case_number = motion_data.get('caseNumber', 'Место для ввода текста.')
    documents = motion_data.get('documents', [])
    
    # Шапка документа с форматированием
    def add_header_paragraph(text):
        p = doc.add_paragraph(text)
        p_format = p.paragraph_format
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)
        p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_format.left_indent = Cm(7)
        return p
    
    add_header_paragraph(f"В {court_name}")
    add_header_paragraph(f"Адрес: {court_address}")
    add_header_paragraph("")
    
    add_header_paragraph(f"Должник:")
    add_header_paragraph(f"{full_name}")
    add_header_paragraph(f"Адрес: {registration.get('address', 'Место для ввода текста.')}")
    add_header_paragraph("")
    
    add_header_paragraph(f"Паспорт: серия {passport.get('series', 'Место для ввода текста.')} номер {passport.get('number', 'Место для ввода текста.')}")
    add_header_paragraph(f"выдан: {passport.get('issuedBy', 'Место для ввода текста.')}")
    add_header_paragraph(f"дата выдачи: {passport.get('issueDate', 'Место для ввода текста.')}")
    add_header_paragraph(f"код подразделения: {passport.get('code', 'Место для ввода текста.')}")
    add_header_paragraph(f"тел. {phone}")
    add_header_paragraph(f"e-mail: {email}")
    add_header_paragraph("")
    add_header_paragraph(f"Номер дела: {case_number}")
    add_header_paragraph("")
    
    doc.add_paragraph()
    
    title = doc.add_heading("ХОДАТАЙСТВО", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_format = title.paragraph_format
    title_format.space_before = Pt(0)
    title_format.space_after = Pt(0)
    title_format.line_spacing = 1.0
    
    subtitle = doc.add_heading("О ПРИОБЩЕНИИ ДОКУМЕНТОВ", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.paragraph_format
    subtitle_format.space_before = Pt(0)
    subtitle_format.space_after = Pt(0)
    subtitle_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # Функция для добавления абзацев с форматированием
    def add_body_paragraph(text):
        p = doc.add_paragraph(text)
        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_format.first_line_indent = Cm(1)
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)
        p_format.line_spacing = 1.0
        return p
    
    # Склоняем ФИО и название суда в родительный падеж
    full_name_genitive = decline_full_name_genitive(full_name)
    court_name_genitive = decline_court_name_genitive(court_name)
    
    # Основной текст
    add_body_paragraph(f"В производстве {court_name_genitive} находится дело № {case_number} по заявлению {full_name_genitive} о признании несостоятельным (банкротом).")
    add_body_paragraph("Согласно статье 41 АПК РФ лица, участвующие в деле, вправе, в том числе, представлять доказательства, заявлять ходатайства.")
    add_body_paragraph("На основании изложенного, Должник просит приобщить к материалам дела следующие документы:")
    
    doc.add_paragraph()
    
    # Список документов
    if documents:
        for idx, doc_title in enumerate(documents, 1):
            add_body_paragraph(f"{idx}. {doc_title}.")
    else:
        add_body_paragraph("1. Место для ввода текста.")
        add_body_paragraph("2. Место для ввода текста.")
        add_body_paragraph("3. Место для ввода текста.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Подпись
    add_body_paragraph(f"Должник: {full_name}")
    add_body_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return base64.b64encode(buffer.read()).decode('utf-8')


def generate_absence_motion_document(
    personal: Dict[str, Any],
    additional: Dict[str, Any],
    motion_data: Dict[str, Any]
) -> str:
    '''Генерирует DOCX документ ходатайства о рассмотрении дела в отсутствие должника'''
    
    doc = Document()
    
    if additional is None:
        additional = {}
    
    passport = personal.get('passport', {})
    registration = personal.get('registration', {})
    
    registration_address = registration.get('address', '')
    auto_court = determine_court_by_address(registration_address)
    
    court_name = additional.get('courtName', auto_court['name'])
    court_address = additional.get('courtAddress', auto_court['address'])
    phone = personal.get('phone', 'Место для ввода текста.')
    email = personal.get('email', 'Место для ввода текста.')
    full_name = personal.get('fullName', 'Место для ввода текста.')
    
    case_number = motion_data.get('caseNumber', 'Место для ввода текста.')
    hearing_date = motion_data.get('hearingDate', '')
    
    # Форматируем дату в русский формат
    if hearing_date:
        from datetime import datetime as dt
        try:
            date_obj = dt.strptime(hearing_date, '%Y-%m-%d')
            formatted_date = date_obj.strftime('%d.%m.%Y')
        except:
            formatted_date = hearing_date
    else:
        formatted_date = 'Место для ввода текста.'
    
    # Шапка документа с форматированием
    def add_header_paragraph(text):
        p = doc.add_paragraph(text)
        p_format = p.paragraph_format
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)
        p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_format.left_indent = Cm(7)
        return p
    
    add_header_paragraph(f"В {court_name}")
    add_header_paragraph(f"Адрес: {court_address}")
    add_header_paragraph("")
    
    add_header_paragraph(f"Должник:")
    add_header_paragraph(f"{full_name}")
    add_header_paragraph(f"Адрес: {registration.get('address', 'Место для ввода текста.')}")
    add_header_paragraph("")
    
    add_header_paragraph(f"Паспорт: серия {passport.get('series', 'Место для ввода текста.')} номер {passport.get('number', 'Место для ввода текста.')}")
    add_header_paragraph(f"выдан: {passport.get('issuedBy', 'Место для ввода текста.')}")
    add_header_paragraph(f"дата выдачи: {passport.get('issueDate', 'Место для ввода текста.')}")
    add_header_paragraph(f"код подразделения: {passport.get('code', 'Место для ввода текста.')}")
    add_header_paragraph(f"тел. {phone}")
    add_header_paragraph(f"e-mail: {email}")
    add_header_paragraph("")
    add_header_paragraph(f"Номер дела: {case_number}")
    add_header_paragraph("")
    
    doc.add_paragraph()
    
    title = doc.add_heading("ХОДАТАЙСТВО", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_format = title.paragraph_format
    title_format.space_before = Pt(0)
    title_format.space_after = Pt(0)
    title_format.line_spacing = 1.0
    
    subtitle = doc.add_heading("О РАССМОТРЕНИИ ДЕЛА В ОТСУТСТВИИ ДОЛЖНИКА", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.paragraph_format
    subtitle_format.space_before = Pt(0)
    subtitle_format.space_after = Pt(0)
    subtitle_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # Функция для добавления абзацев с форматированием
    def add_body_paragraph(text):
        p = doc.add_paragraph(text)
        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_format.first_line_indent = Cm(1)
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)
        p_format.line_spacing = 1.0
        return p
    
    # Склоняем ФИО и название суда в родительный падеж
    full_name_genitive = decline_full_name_genitive(full_name)
    court_name_genitive = decline_court_name_genitive(court_name)
    
    # Основной текст
    add_body_paragraph(f"В производстве {court_name_genitive} находится дело № {case_number} по заявлению {full_name_genitive} о признании несостоятельным (банкротом).")
    add_body_paragraph(f"{formatted_date} г. назначено судебное заседание по рассмотрению указанного дела.")
    add_body_paragraph("В силу ч. 2 ст. 156 Арбитражного процессуального кодекса Российской Федерации стороны вправе известить арбитражный суд о возможности рассмотрения дела в их отсутствие.")
    add_body_paragraph("На основании изложенного, руководствуясь ч. 2 ст. 156, ч. 1 ст. 41 Арбитражного процессуального кодекса Российской Федерации,")
    
    doc.add_paragraph()
    
    # Заголовок ПРОШУ
    prosh_heading = doc.add_paragraph("ПРОШУ:")
    prosh_heading.runs[0].bold = True
    prosh_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    add_body_paragraph(f"рассмотреть дело № {case_number} по существу в отсутствие заявителя.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Подпись для absence motion
    add_body_paragraph(f"Должник: {full_name}")
    add_body_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
    
    buffer_absence = io.BytesIO()
    doc.save(buffer_absence)
    buffer_absence.seek(0)
    return base64.b64encode(buffer_absence.read()).decode('utf-8')


def generate_property_exclusion_motion_document(
    personal: Dict[str, Any],
    additional: Dict[str, Any],
    property_data: Dict[str, Any],
    motion_data: Dict[str, Any]
) -> str:
    '''Генерирует DOCX документ ходатайства об исключении имущества из конкурсной массы'''
    
    doc = Document()
    
    if additional is None:
        additional = {}
    
    passport = personal.get('passport', {})
    registration = personal.get('registration', {})
    
    registration_address = registration.get('address', '')
    auto_court = determine_court_by_address(registration_address)
    
    court_name = additional.get('courtName', auto_court['name'])
    court_address = additional.get('courtAddress', auto_court['address'])
    phone = personal.get('phone', 'Место для ввода текста.')
    email = personal.get('email', 'Место для ввода текста.')
    full_name = personal.get('fullName', 'Место для ввода текста.')
    
    case_number = motion_data.get('caseNumber', 'Место для ввода текста.')
    
    # Находим единственное жилье
    sole_residence = None
    if property_data and property_data.get('realEstate'):
        for item in property_data.get('realEstate', []):
            if item.get('isSoleResidence', False):
                sole_residence = item
                break
    
    # Шапка документа с форматированием
    def add_header_paragraph(text):
        p = doc.add_paragraph(text)
        p_format = p.paragraph_format
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)
        p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_format.left_indent = Cm(7)
        return p
    
    add_header_paragraph(f"В {court_name}")
    add_header_paragraph(f"Адрес: {court_address}")
    add_header_paragraph("")
    
    add_header_paragraph(f"Должник:")
    add_header_paragraph(f"{full_name}")
    add_header_paragraph(f"Адрес: {registration.get('address', 'Место для ввода текста.')}")
    add_header_paragraph("")
    
    add_header_paragraph(f"Паспорт: серия {passport.get('series', 'Место для ввода текста.')} номер {passport.get('number', 'Место для ввода текста.')}")
    add_header_paragraph(f"выдан: {passport.get('issuedBy', 'Место для ввода текста.')}")
    add_header_paragraph(f"дата выдачи: {passport.get('issueDate', 'Место для ввода текста.')}")
    add_header_paragraph(f"код подразделения: {passport.get('code', 'Место для ввода текста.')}")
    add_header_paragraph(f"тел. {phone}")
    add_header_paragraph(f"e-mail: {email}")
    add_header_paragraph("")
    add_header_paragraph(f"Номер дела: {case_number}")
    add_header_paragraph("")
    
    doc.add_paragraph()
    
    title = doc.add_heading("ХОДАТАЙСТВО", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_format = title.paragraph_format
    title_format.space_before = Pt(0)
    title_format.space_after = Pt(0)
    title_format.line_spacing = 1.0
    
    subtitle = doc.add_heading("ОБ ИСКЛЮЧЕНИИ ИМУЩЕСТВА ИЗ КОНКУРСНОЙ МАССЫ", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.paragraph_format
    subtitle_format.space_before = Pt(0)
    subtitle_format.space_after = Pt(0)
    subtitle_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # Функция для добавления абзацев с форматированием
    def add_body_paragraph(text):
        p = doc.add_paragraph(text)
        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_format.first_line_indent = Cm(1)
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)
        p_format.line_spacing = 1.0
        return p
    
    # Склоняем ФИО и название суда в родительный падеж
    full_name_genitive = decline_full_name_genitive(full_name)
    court_name_genitive = decline_court_name_genitive(court_name)
    
    # Основной текст
    add_body_paragraph(f"В производстве {court_name_genitive} находится дело № {case_number} по заявлению {full_name_genitive} о признании несостоятельным (банкротом).")
    
    doc.add_paragraph()
    
    add_body_paragraph("В соответствии со ст. 213.25 Федерального закона от 26.10.2002 N 127-ФЗ «О несостоятельности (банкротстве)» (далее - Закон о банкротстве) все имущество гражданина, имеющееся на дату принятия решения арбитражного суда о признании гражданина банкротом и введении реализации имущества гражданина и выявленное или приобретенное после даты принятия указанного решения, составляет конкурсную массу, за исключением имущества, определенного п. 3 названной статьи.")
    add_body_paragraph("Согласно п. 3 ст. 213.25 Закона о банкротстве из конкурсной массы исключается имущество, на которое не может быть обращено взыскание в соответствии с гражданским процессуальным законодательством.")
    
    doc.add_paragraph()
    
    add_body_paragraph("Частью 1 ст. 446 ГПК РФ предусмотрено, что взыскание по исполнительным документам не может быть обращено на жилое помещение (его части), принадлежащее гражданину-должнику на праве собственности, если для гражданина-должника и членов его семьи, совместно проживающих в принадлежащем помещении, оно является единственным пригодным для постоянного проживания помещением, за исключением указанного в настоящем абзаце имущества, если оно является предметом ипотеки и на него в соответствии с законодательством об ипотеке может быть обращено взыскание.")
    add_body_paragraph("Согласно ч. 1 ст. 16 ЖК РФ к жилым помещениям относятся: жилой дом, часть жилого дома; квартира, часть квартиры; комната.")
    add_body_paragraph("По смыслу приведенных правовых норм, предусмотренный абз. 2 ч. 1 ст. 446 ГПК РФ имущественный (исполнительский) иммунитет означает запрет обращать взыскание на любой вид жилого помещения, определенный Жилищным кодексом Российской Федерации, если соответствующее жилое помещение является для должника и членов его семьи единственным пригодным для постоянного проживания помещением и не является предметом ипотеки.")
    
    doc.add_paragraph()
    
    # Описание имущества
    if sole_residence:
        area_text = str(sole_residence.get('area', ''))
        land_area = sole_residence.get('landArea')
        land_text = ""
        if land_area and land_area != 0:
            land_text = f" с земельным участком площадью {land_area} кв. м."
        
        property_description = f"У должника в собственности имеется {sole_residence.get('type', 'недвижимость')}, общей площадью {area_text} кв. м.{land_text} по адресу: {sole_residence.get('address', 'Место для ввода текста.')}, которое является единственным жильем для должника и членов его семьи. Данный объект недвижимости не является предметом ипотеки."
    else:
        property_description = "У должника в собственности имеется жилое помещение, которое является единственным жильем для должника и членов его семьи. Данный объект недвижимости не является предметом ипотеки."
    
    add_body_paragraph(property_description)
    add_body_paragraph("На основании вышеизложенного данное имущество подлежит исключению из конкурсной массы.")
    
    doc.add_paragraph()
    
    add_body_paragraph("На основании вышеизложенного и, руководствуясь статьей 213.25 Федерального закона от 26.10.2002 N 127-ФЗ «О несостоятельности (банкротстве)», прошу:")
    
    doc.add_paragraph()
    
    # Заголовок ПРОШУ (уже включен в текст выше, но для структуры можно оставить отдельный параграф)
    # Просительная часть
    if sole_residence:
        area_text = str(sole_residence.get('area', ''))
        land_area = sole_residence.get('landArea')
        land_text = ""
        if land_area and land_area != 0:
            land_text = f" с земельным участком площадью {land_area} кв. м."
        
        petition_text = f"1. Исключить из конкурсной массы {full_name_genitive} следующее имущество: {sole_residence.get('type', 'недвижимость')}, общей площадью {area_text} кв. м.{land_text} по адресу: {sole_residence.get('address', 'Место для ввода текста.')}."
    else:
        petition_text = f"1. Исключить из конкурсной массы {full_name_genitive} следующее имущество: жилое помещение, являющееся единственным пригодным для постоянного проживания."
    
    add_body_paragraph(petition_text)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Подпись
    add_body_paragraph(f"Должник: {full_name}")
    add_body_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
    
    buffer_property = io.BytesIO()
    doc.save(buffer_property)
    buffer_property.seek(0)
    return base64.b64encode(buffer_property.read()).decode('utf-8')


def generate_debt_discharge_motion_document(
    personal: Dict[str, Any],
    additional: Dict[str, Any],
    children_data: Dict[str, Any],
    motion_data: Dict[str, Any]
) -> str:
    '''Генерирует DOCX документ ходатайства об освобождении должника от долгов'''
    
    doc = Document()
    
    if additional is None:
        additional = {}
    
    passport = personal.get('passport', {})
    registration = personal.get('registration', {})
    
    registration_address = registration.get('address', '')
    auto_court = determine_court_by_address(registration_address)
    
    court_name = additional.get('courtName', auto_court['name'])
    court_address = additional.get('courtAddress', auto_court['address'])
    phone = personal.get('phone', 'Место для ввода текста.')
    email = personal.get('email', 'Место для ввода текста.')
    full_name = personal.get('fullName', 'Место для ввода текста.')
    
    case_number = motion_data.get('caseNumber', 'Место для ввода текста.')
    hearing_date = motion_data.get('hearingDate', '')
    property_status = motion_data.get('propertyStatus', 'not-found')
    property_sale_amount = motion_data.get('propertySaleAmount', '')
    no_contestable_transactions = motion_data.get('noContestablTransactions', True)
    has_employment = motion_data.get('hasEmployment', False)
    employer_name = motion_data.get('employerName', '')
    monthly_income = motion_data.get('monthlyIncome', '')
    total_debt = motion_data.get('totalDebt', '0')
    
    # Форматируем дату в русский формат
    if hearing_date:
        from datetime import datetime as dt
        try:
            date_obj = dt.strptime(hearing_date, '%Y-%m-%d')
            formatted_date = date_obj.strftime('%d.%m.%Y')
        except:
            formatted_date = hearing_date
    else:
        formatted_date = 'Место для ввода текста.'
    
    # Шапка документа с форматированием
    def add_header_paragraph(text):
        p = doc.add_paragraph(text)
        p_format = p.paragraph_format
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)
        p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_format.left_indent = Cm(7)
        return p
    
    add_header_paragraph(f"В {court_name}")
    add_header_paragraph(f"Адрес: {court_address}")
    add_header_paragraph("")
    
    add_header_paragraph(f"Должник:")
    add_header_paragraph(f"{full_name}")
    add_header_paragraph(f"Адрес: {registration.get('address', 'Место для ввода текста.')}")
    add_header_paragraph("")
    
    add_header_paragraph(f"Паспорт: серия {passport.get('series', 'Место для ввода текста.')} номер {passport.get('number', 'Место для ввода текста.')}")
    add_header_paragraph(f"выдан: {passport.get('issuedBy', 'Место для ввода текста.')}")
    add_header_paragraph(f"дата выдачи: {passport.get('issueDate', 'Место для ввода текста.')}")
    add_header_paragraph(f"код подразделения: {passport.get('code', 'Место для ввода текста.')}")
    add_header_paragraph(f"тел. {phone}")
    add_header_paragraph(f"e-mail: {email}")
    add_header_paragraph("")
    add_header_paragraph(f"Номер дела: {case_number}")
    add_header_paragraph("")
    
    doc.add_paragraph()
    
    title = doc.add_heading("ХОДАТАЙСТВО", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_format = title.paragraph_format
    title_format.space_before = Pt(0)
    title_format.space_after = Pt(0)
    title_format.line_spacing = 1.0
    
    subtitle = doc.add_heading("ОБ ОСВОБОЖДЕНИИ ДОЛЖНИКА ОТ ДОЛГОВ", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.paragraph_format
    subtitle_format.space_before = Pt(0)
    subtitle_format.space_after = Pt(0)
    subtitle_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # Функция для добавления абзацев с форматированием
    def add_body_paragraph(text):
        p = doc.add_paragraph(text)
        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_format.first_line_indent = Cm(1)
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(0)
        p_format.line_spacing = 1.0
        return p
    
    # Склоняем ФИО в разные падежи
    full_name_genitive = decline_full_name_genitive(full_name)
    full_name_accusative = decline_full_name_accusative(full_name)
    court_name_genitive = decline_court_name_genitive(court_name)
    
    # Вводный абзац с номером дела
    add_body_paragraph(f"В производстве {court_name_genitive} находится дело № {case_number} по заявлению {full_name_genitive} о признании несостоятельным (банкротом).")
    
    doc.add_paragraph()
    
    # Основной текст
    add_body_paragraph(f"Судебное заседание по рассмотрению отчета финансового управляющего назначено на {formatted_date} г.")
    add_body_paragraph("Настоящим, Должник сообщает Суду следующую информацию по результату проведенных мероприятий:")
    
    doc.add_paragraph()
    
    # Пункт об имуществе
    if property_status == 'not-found':
        add_body_paragraph("• В ходе всех предусмотренных законом проведенных мероприятий, имущества, подлежащего включению в конкурсную массу не выявлено.")
    else:
        sale_amount_formatted = format_number(property_sale_amount) if property_sale_amount else 'Место для ввода текста.'
        add_body_paragraph(f"• В ходе всех предусмотренных законом проведенных мероприятий имущество выявлено, добровольно передано финансовому управляющему и реализовано на торгах на сумму {sale_amount_formatted} рублей.")
    
    # Пункт о сделках (условно)
    if no_contestable_transactions:
        add_body_paragraph("• Сделок должника, подлежащих оспариванию не выявлено.")
    
    # Признаки фиктивного банкротства
    add_body_paragraph("• Признаков фиктивного/преднамеренного банкротства не выявлено.")
    
    # Лица на иждивении
    has_minor_children = children_data and not children_data.get('noChildren', True) and children_data.get('children')
    if has_minor_children:
        children_list = children_data.get('children', [])
        children_names = ', '.join([child.get('fullName', 'ФИО') for child in children_list])
        add_body_paragraph(f"• У должника имеются несовершеннолетние дети на иждивении: {children_names}.")
    else:
        add_body_paragraph("• Лиц на иждивении не имеется.")
    
    # Трудовая деятельность
    if has_employment:
        income_formatted = format_number(monthly_income) if monthly_income else 'Место для ввода текста.'
        add_body_paragraph(f"• Должник в настоящее время осуществляет трудовую деятельность в {employer_name} и получает доход в размере {income_formatted} рублей, что не хватает для расчетов с кредиторами.")
    else:
        add_body_paragraph("• Должник в настоящее время трудовую деятельность не осуществляет.")
    
    # Реестр требований кредиторов
    debt_formatted = format_number(total_debt) if total_debt else 'Место для ввода текста.'
    add_body_paragraph(f"• Реестр требований кредиторов сформирован в общей сумме {debt_formatted} рублей.")
    
    doc.add_paragraph()
    
    # Правовое обоснование
    add_body_paragraph("Как неоднократно отмечал Верховный Суд Российской Федерации, основной задачей института потребительского банкротства является социальная реабилитация гражданина - предоставление ему возможности заново выстроить экономические отношения, законно избавившись от необходимости отвечать по старым обязательствам, что в определенной степени ущемляет права кредиторов должника, вследствие чего к должнику законодателем предъявляются повышенные требования в части добросовестности, подразумевающие, помимо прочего, честное сотрудничество с финансовым управляющим и кредиторами, открытое взаимодействие с судом.")
    add_body_paragraph("Таким образом, такое положение гражданина-банкрота перед кредиторами возможно при условии его добросовестного поведения с кредиторами (в том числе до возбуждения дела о банкротстве), управляющим и судом.")
    
    doc.add_paragraph()
    
    add_body_paragraph("В настоящем деле о банкротстве при проведении всех мероприятий, предусмотренных Законом в рамках процедуры банкротства гражданина, со стороны Должника не было установлено недобросовестного поведения, не были выявлены сделки, совершенные Должником во вред кредиторам, не установлены факты сокрытия информации и пр., таким образом отсутствуют основания о не списании долгов и Должник ходатайствует о завершении процедуры банкротства в его отношении и полном списании всех долгов.")
    
    doc.add_paragraph()
    
    add_body_paragraph("На основании изложенного,")
    
    doc.add_paragraph()
    
    # Заголовок ПРОШУ
    prosh_heading = doc.add_paragraph("ПРОШУ:")
    prosh_heading.runs[0].bold = True
    prosh_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    add_body_paragraph(f"1. Завершить процедуру банкротства в отношении {full_name_genitive}.")
    add_body_paragraph(f"2. Освободить {full_name_accusative} от дальнейшего исполнения обязанностей перед кредиторами включенными в реестр требований кредиторов, а также тех, которые имелись у Должника на дату введения процедуры банкротства.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Подпись
    add_body_paragraph(f"Должник: {full_name}")
    add_body_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return base64.b64encode(buffer.read()).decode('utf-8')